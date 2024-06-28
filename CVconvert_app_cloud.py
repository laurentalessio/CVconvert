import streamlit as st
import PyPDF2
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openai import OpenAI
import io

def read_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def read_docx(file):
    doc = Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def process_cv(consultant_cv, template_cv, api_key):
    client = OpenAI(api_key=api_key)
    try:
        prompt = f"""
        You are tasked with formatting a consultant's CV according to the THREE60 template structure.
        Use ONLY the information from the consultant's CV. DO NOT use any specific information from the template CV.
        
        The structure should be as follows:
        [HEADER]
        Consultant's Name
        Consultant's Current Position
        [/HEADER]

        [SECTION]Years of experience
        Content
        [/SECTION]

        [SECTION]Discipline
        Content
        [/SECTION]

        [SECTION]Role
        Content
        [/SECTION]

        [SECTION]Technical skills
        Content
        [/SECTION]

        [SECTION]Professional skills
        Content
        [/SECTION]

        [SECTION]Professional Summary
        Content
        [/SECTION]

        [SECTION]Work Experience - Summary
        Content
        [/SECTION]

        [SECTION]Work Experience - Detailed
        Content
        [/SECTION]

        [SECTION]Education and training
        Content
        [/SECTION]

        [SECTION]Personal skills and competencies
        Content
        [/SECTION]

        Fill each section with relevant information from the consultant's CV. If information for a section is not available, write 'Information not provided' as the content.
        Do not invent or assume any information not present in the consultant's CV.

        Consultant CV to format:
        {consultant_cv}

        Please provide the formatted CV content, using ONLY information from the consultant's CV.
        """

        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a CV formatting assistant. Only use information provided in the consultant's CV."},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error processing CV: {str(e)}")
        return None

def create_word_document(formatted_cv, template_file):
    # Load the template
    doc = Document(template_file)
    
    # Parse the formatted CV
    header = formatted_cv.split('[/HEADER]')[0].replace('[HEADER]', '').strip().split('\n')
    sections = formatted_cv.split('[SECTION]')[1:]
    
    # Update the header (usually in the first table)
    if doc.tables:
        table = doc.tables[0]
        if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
            cell = table.rows[0].cells[0]
            cell.text = '\n'.join(header)

    # Process each section
    for section in sections:
        lines = section.split('\n')
        if lines:
            heading = lines[0].strip()
            content = '\n'.join(lines[1:]).strip()
            
            # Find existing paragraph with this heading
            for paragraph in doc.paragraphs:
                if paragraph.text.strip() == heading:
                    # Clear existing content
                    paragraph.clear()
                    # Add new content
                    paragraph.add_run(heading).bold = True
                    if content:
                        paragraph.add_run('\n' + content)
                    break
            else:
                # If heading not found, create new paragraph
                new_para = doc.add_paragraph()
                new_para.add_run(heading).bold = True
                if content:
                    new_para.add_run('\n' + content)
    
    # Save the document to a bytes buffer
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

def main():
    st.title("CV Formatter for Three60 Template")

    # Add input for OpenAI API key
    api_key = st.text_input("Enter your OpenAI API key", type="password")

    consultant_cv_file = st.file_uploader("Upload Consultant CV (PDF)", type="pdf")
    template_cv_file = st.file_uploader("Upload Three60 Template Example (DOCX)", type="docx")

    if consultant_cv_file and template_cv_file and api_key:
        consultant_cv = read_pdf(consultant_cv_file)
        template_cv = read_docx(template_cv_file)

        if st.button("Process CV"):
            formatted_cv = process_cv(consultant_cv, template_cv, api_key)
            if formatted_cv:
                doc_buffer = create_word_document(formatted_cv, template_cv_file)
                st.success("CV formatted successfully!")
                st.download_button(
                    label="Download Formatted CV",
                    data=doc_buffer,
                    file_name="formatted_cv.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    else:
        st.warning("Please upload both CV files and enter your OpenAI API key to proceed.")

if __name__ == "__main__":
    main()
